import json
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_custom_style(doc, style_name, font_name, font_size, bold=False, color=None):
    styles = doc.styles
    if style_name not in styles:
        style = styles.add_style(style_name, 1)  # 1 for paragraph style
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = bold
        if color:
            font.color.rgb = RGBColor(color[0], color[1], color[2])
    return styles[style_name]

def set_cell_color(cell, color):
    """Set background color of a cell."""
    cell_properties = cell._element.get_or_add_tcPr()
    cell_shading = OxmlElement('w:shd')
    cell_shading.set(qn('w:fill'), color)
    cell_properties.append(cell_shading)

#Get the directory of the Flask application file
app_dir = os.path.dirname(__file__)

# JSON file Path
json_path = os.path.join(app_dir, "result", "json")

# Find the JSON file in the directory
json_file_path = None
for filename in os.listdir(json_path):
    if filename.endswith(".json"):
        json_file_path = os.path.join(json_path, filename)
        break

# Check if a JSON file was found
if json_file_path is None:
    print("No JSON file found in the directory.")
    exit()
# Load JSON data from file
with open(json_file_path, 'r') as file:
    json_data = json.load(file)

# Create a new Document
doc = Document()

# Add heading
heading = doc.add_heading(level=1)
run = heading.add_run('Appendix A.1 Certificate Details')
run.bold = True
run.font.size = Pt(14)
heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


# Add a table
table = doc.add_table(rows=0, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# Add header row with merged cells for title
header_cells = table.add_row().cells
header_cells[0].merge(header_cells[1])
header_run = header_cells[0].paragraphs[0].add_run('Certificate')
header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
header_run.bold = True
header_run.font.size = Pt(12)
header_run.font.color.rgb = RGBColor(255, 255, 255)
set_cell_color(header_cells[0], '4F81BD')

# Find the 'ip' value for 'cert_commonName'
common_name_item = next((item for item in json_data if item['id'] == 'cert_commonName'), None)
if common_name_item:
    ip_value = common_name_item['ip']
    port_value = common_name_item['port']
    ip_port_value = f"{ip_value}:{port_value}"
else:
    ip_port_value = 'Not Available'

# Add second header row with the dynamically found IP value
url_cells = table.add_row().cells
url_cells[0].merge(url_cells[1])
url_run = url_cells[0].paragraphs[0].add_run(ip_port_value)
url_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
url_run.bold = True
url_run.font.size = Pt(12)
url_run.font.color.rgb = RGBColor(255, 255, 255)
set_cell_color(url_cells[0], '4F81BD')

# Define the row color mappings for severity
severity_color_map = {
    "OK": "D9EAD3",  # Green
    "INFO": "FFFFFF",  # Plain White
    "LOW": "FFCC99",  # Unknown Color
    "MEDIUM": "FFCC33",  # Yellow
    "WARN": "FFCC33",  # Yellow
    "CRITICAL": "F4CCCC"  # Red
}

# Define the mappings of the keys to search in the JSON file
key_to_id_map = {
    "Serial Number": "cert_serialNumber",
    "Subject": "cert_commonName",
    "Issuer": "cert_caIssuers",
    "Certificate Chain": "cert_trust",
    "Valid From": "cert_notBefore",
    "Valid Until": "cert_notAfter",
    "Signature Algorithm": "cert_signatureAlgorithm",
    "Public Key Size": "cert_keySize",
    "DNS Certificate Authority Authorisation Record": "DNS_CAArecord",
    #Amend Host Name Validation
    "Hostname Validation": "cert_hostNameValidation",
    #Amend Self-signed
    "Self-Signed": "cert_selfSigned",
    "OCSP URL": "cert_ocspURL",
    "Subject Alternative Name": "cert_subjectAltName",
    "OCSP Stapling": "OCSP_stapling",
    "OCSP Must Staple Extension": "cert_mustStapleExtension"
}

# Add the JSON data to the table
for idx, (key, json_id) in enumerate(key_to_id_map.items()):
    row_cells = table.add_row().cells
    key_cell = row_cells[0]
    value_cell = row_cells[1]

    key_run = key_cell.paragraphs[0].add_run(key)
    key_run.font.size = Pt(10)
    key_run.bold = True
    
    # Find the corresponding value and severity for the given id in the JSON data
    item = next((item for item in json_data if item['id'] == json_id), None)
    if item:
        value = item['finding']
        severity = item['severity']
        # Replacing -- value with N/A
        if value == '--':
            value = 'N/A'
    else:
        value = 'Not Available'
        severity = 'INFO'  # Default to INFO if not found

    value_run = value_cell.paragraphs[0].add_run(value)
    value_run.font.size = Pt(10)

    # Set background color based on severity
    bg_color = severity_color_map.get(severity, "FFFFFF")
    set_cell_color(value_cell, bg_color)

# Add heading for TLS/SSL Protocol Support
heading = doc.add_heading(level=1)
run = heading.add_run('Appendix A.2 TLS/SSL Protocol Support')
run.bold = True
run.font.size = Pt(14)
heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Define the mappings for TLS/SSL Protocol Support
tls_protocols = ["TLS 1.3", "TLS 1.2", "TLS 1.1", "TLS 1.0", "SSL 3.0", "SSL 2.0"]
protocol_id_map = {
    "TLS 1.3": "TLS1_3",
    "TLS 1.2": "TLS1_2",
    "TLS 1.1": "TLS1_1",
    "TLS 1.0": "TLS1",
    "SSL 3.0": "SSLv3",
    "SSL 2.0": "SSLv2"
}

# Add a table for TLS/SSL Protocol Support
table = doc.add_table(rows=0, cols=len(tls_protocols) + 1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# Add header row for protocols
header_cells = table.add_row().cells
header_cells[0].paragraphs[0].add_run(ip_port_value).bold = True
for i, protocol in enumerate(tls_protocols, 1):
    header_run = header_cells[i].paragraphs[0].add_run(protocol)
    header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run.bold = True
    header_run.font.size = Pt(12)
    header_run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_color(header_cells[i], '4F81BD')

# Add data row for protocol support
data_cells = table.add_row().cells
data_cells[0]._element.get_or_add_tcPr().append(OxmlElement('w:vMerge'))
for i, protocol in enumerate(tls_protocols, 1):
    item = next((item for item in json_data if item['id'] == protocol_id_map[protocol]), None)
    if item:
        finding = item['finding']
        severity = item['severity']
        if finding == "offered" or finding == "offered with final":
            value = "Yes"
        else:
            value = "No"
        # Set background color based on severity
        bg_color = severity_color_map.get(severity, "FFFFFF")
    else:
        value = "Not Available"
        bg_color = "FFFFFF"
    cell_run = data_cells[i].paragraphs[0].add_run(value)
    cell_run.font.size = Pt(10)
    set_cell_color(data_cells[i], bg_color)

# Add heading for TLS/SSL Protocol Issues
heading = doc.add_heading(level=1)
run = heading.add_run('Appendix A.3 TLS/SSL Protocol Issues')
run.bold = True
run.font.size = Pt(14)
heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Define the mappings for TLS/SSL Protocol Issues
tls_protocol_issues = ["Client-Initiated Renegotation",
                 "Secure Renegotiation", 
                 "OpenSSL CCS Injection(CVE-2014-0224)",
                 "POODLE (SSL 3) (CVE-2014-3566)",
                 "ROBOT", "BREACH (CVE-2013-3587)","HTTP Strict Transport Security (HSTS)","HSTS PreLoad"]
issue_id_map = {
    #Amend Client-Initiated Renegotation
    "Client-Initiated Renegotation": "TLS1_3",
    "Secure Renegotiation": "secure_renego",
    "OpenSSL CCS Injection(CVE-2014-0224)": "CCS",
    "POODLE (SSL 3) (CVE-2014-3566)": "POODLE_SSL",
    "ROBOT": "ROBOT",
    "BREACH (CVE-2013-3587)": "BREACH",
    # Amend HSTS 
    "HTTP Strict Transport Security (HSTS)": "RC4",
    "HSTS PreLoad": "HSTS_preload"
}

# Add a table for TLS/SSL Protocol Support
table = doc.add_table(rows=0, cols=len(tls_protocol_issues) + 1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# Add header row for protocols
header_cells = table.add_row().cells
header_cells[0].paragraphs[0].add_run(ip_port_value).bold = True
for i, issues in enumerate(tls_protocol_issues, 1):
    header_run = header_cells[i].paragraphs[0].add_run(issues)
    header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run.bold = True
    header_run.font.size = Pt(12)
    header_run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_color(header_cells[i], '4F81BD')

# Add data row for protocol support
data_cells = table.add_row().cells
data_cells[0]._element.get_or_add_tcPr().append(OxmlElement('w:vMerge'))
for i, issues in enumerate(tls_protocol_issues, 1):
    item = next((item for item in json_data if item['id'] == issue_id_map[issues]), None)
    if item:
        finding = item['finding'].split(',')
        severity = item['severity']
        # Determine the value based on the findings
        if any(finding.strip() in ["not supported", "vulnerable","not vulnerable, no gzip/deflate/compress/br HTTP compression  - only supplied '/' tested"] for finding in findings):
            value = "Yes"
        else:
            value = "No"
        # Set background color based on severity
        bg_color = severity_color_map.get(severity, "FFFFFF")
    else:
        value = "Not Available"
        bg_color = "FFFFFF"
    cell_run = data_cells[i].paragraphs[0].add_run(value)
    cell_run.font.size = Pt(10)
    set_cell_color(data_cells[i], bg_color)

# Save the document
doc.save('Certificate_Details_Report.docx')

