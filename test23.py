import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_custom_style(doc, style_name, font_name, font_size, bold=False, color=None):
    styles = doc.styles
    if style_name not in styles:
        style = styles.add_style(style_name, 1)  # 1 for paragraph style
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = bold
        if color:
            font.color.rgb = RGBColor(color[0], color[1], color[2])
    return styles[style_name]

def set_cell_color(cell, color):
    """Set background color of a cell."""
    cell_properties = cell._element.get_or_add_tcPr()
    cell_shading = OxmlElement('w:shd')
    cell_shading.set(qn('w:fill'), color)
    cell_properties.append(cell_shading)

# Load JSON data from file
json_file_path = '/home/ubuntu/TestSSL/testssl.sh-webfrontend/result/json/interceptica.com_p443-20240503-0552.json'
with open(json_file_path, 'r') as file:
    json_data = json.load(file)

# Create a new Document
doc = Document()

# Add heading
heading = doc.add_heading(level=1)
run = heading.add_run('Appendix A.1 Certificate Details')
run.bold = True
run.font.size = Pt(14)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add a table
table = doc.add_table(rows=0, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# Add header row with merged cells for title
header_cells = table.add_row().cells
header_cells[0].merge(header_cells[1])
header_run = header_cells[0].paragraphs[0].add_run('Certificate')
header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
header_run.bold = True
header_run.font.size = Pt(12)
header_run.font.color.rgb = RGBColor(255, 255, 255)
set_cell_color(header_cells[0], '4F81BD')

# Add second header row with URL
url_cells = table.add_row().cells
url_cells[0].merge(url_cells[1])
url_run = url_cells[0].paragraphs[0].add_run('www.interceptica.com:443 (34.149.87.45)')
url_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
url_run.bold = True
url_run.font.size = Pt(12)
url_run.font.color.rgb = RGBColor(255, 255, 255)
set_cell_color(url_cells[0], '95B3D7')

# Define the row color mappings (alternative colors)
row_colors = ['D9EAD3', 'EAD1DC', 'F4CCCC', 'FFF2CC', 'C9DAF8']

# Define the mappings of the keys to search in the JSON file
key_to_id_map = {
    "Serial Number": "cert_serialNumber",
    "Subject": "cert_commonName",
    "Issuer": "cert_issuer",
    "Certificate Chain": "cert_chain_of_trust",
    "Valid From": "cert_validFrom",
    "Valid Until": "cert_validTo",
    "Signature Algorithm": "cert_signatureAlgorithm",
    "Public Key Size": "cert_keySize",
    "DNS Certificate Authority Authorisation Record": "cert_dane",
    "Hostname Validation": "cert_hostNameValidation",
    "Self-Signed": "cert_selfSigned"
}

# Add the JSON data to the table
for idx, (key, json_id) in enumerate(key_to_id_map.items()):
    row_cells = table.add_row().cells
    key_cell = row_cells[0]
    value_cell = row_cells[1]

    key_run = key_cell.paragraphs[0].add_run(key)
    key_run.font.size = Pt(10)
    key_run.bold = True
    
    # Find the corresponding value for the given id in the JSON data
    value = next((item['finding'] for item in json_data if item['id'] == json_id), 'Not Available')
    value_run = value_cell.paragraphs[0].add_run(value)
    value_run.font.size = Pt(10)

    # Set background color for each row alternatively
    set_cell_color(key_cell, row_colors[idx % len(row_colors)])
    set_cell_color(value_cell, row_colors[idx % len(row_colors)])

# Save the document
doc.save('Certificate_Details_Report.docx')
